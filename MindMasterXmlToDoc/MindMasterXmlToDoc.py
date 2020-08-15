#!/usr/bin/python
# -*- coding: UTF-8 -*-
  
import xml.sax
import sys

from docx import Document
from docx.shared import Inches

#递归报错的处理
sys.setrecursionlimit(9000000) 
 
key=["Shape","Text","LevelData"]


class XmlHandler( xml.sax.ContentHandler ):
   def __init__(self):
      self.CurrentData = []
      self.shape={}
      self.type=[]
      self.shapes=[]
 
   # 元素开始事件处理
   def startElement(self, tag, attributes):
        if(tag=="Shape"):
            shapeType = attributes.getValue("Type")
            if shapeType not in self.type:
                self.type.append(shapeType)
            id = attributes.getValue("ID")    
            self.shape["id"] = id 
            self.shape["type"] = shapeType
        if(tag=="Super"):
            superLevel = attributes.getValue("V")
            self.shape["super"] = superLevel 
        if(len(self.CurrentData)==0):
           self.CurrentData.append(tag) 
           print(self.CurrentData[0])
           return 

        self.CurrentData.append(tag)
        if tag in key or "LevelData" in self.CurrentData:
            a = "" 
            if attributes.__len__() > 0:#attrs.getLength()>0
                for _,attr in enumerate(attributes.getNames()):
                    a = a + " " + attr + "=" +  attributes.getValue(attr)
            self.printTag(tag +" " + a,len(self.CurrentData))




   def printTag(self,tag, l):
       i = 0 
       while i < l :
           print("  ",end="")
           i += 1
       print(tag)
       return

   # 元素结束事件处理
   def endElement(self, tag):
        if tag in key or "LevelData" in self.CurrentData:
            #self.printTag("/" +tag,len(self.CurrentData))
            pass
        del self.CurrentData[len(self.CurrentData)-1] 
        if(tag == "Shape" and self.shape.__contains__("content")):
            self.shapes.append(self.shape)
            self.shape={}    
        


   # 内容事件处理
   def characters(self, content):
       #tag = self.CurrentData[len(self.CurrentData) -1]
       #if tag in key or "LevelData" in self.CurrentData:
           self.printTag("content" + content,len(self.CurrentData) + 1)
           if(self.shape.__contains__("content")):
               self.shape["content"] = self.shape["content"] + content
           else:
                self.shape["content"] = content
       #pass
   def find(self,parentShape,shapes):
        if(parentShape["id"] == "367"):
           print(parentShape)
        for shape in shapes:
            if(shape["super"] == parentShape["id"]):
                if(not parentShape.__contains__("sub")):
                    parentShape["sub"] = []                
                parentShape["sub"].append(shape)
                #shapes.remove(shape)
                self.find(shape,shapes)

   def exportToDoc(self,exportShape):
        # 创建word文档对象
        document = Document()
        # 添加标题
        document.add_heading(exportShape["content"], 0)
        if(exportShape.__contains__("sub")):
            for shape in exportShape["sub"]:
                self.exportDoc(document,shape,1)
        document.save("D:\jyw\study\serverless\page\page.docx")             


   def exportDoc(self,document,exportShape,index):
        document.add_heading(exportShape["content"], index)
        if(exportShape.__contains__("sub")):
            for shape in exportShape["sub"]:
                self.exportDoc(document,shape,index + 1)        

if ( __name__ == "__main__"):
   
   # 创建一个 XMLReader
   parser = xml.sax.make_parser()
   # turn off namepsaces
   parser.setFeature(xml.sax.handler.feature_namespaces, 0)
 
   # 重写 ContextHandler
   Handler = XmlHandler()
   parser.setContentHandler( Handler )
   
   parser.parse("D:\jyw\study\serverless\page\page.xml")
   print(Handler.type)
   for shape in Handler.shapes:
       print(shape)


   for shape in Handler.shapes:
       if(shape["type"] == "MainIdea"):
           parentShape = shape
           Handler.shapes.remove(shape)
           break
   Handler.find(parentShape,Handler.shapes) 
   print(parentShape)
   Handler.exportToDoc(parentShape)

